import json
import chardet
import numpy as np
import re
from Levenshtein import distance as lev_distance
from nltk.stem import SnowballStemmer
from Sastrawi.Stemmer.StemmerFactory import StemmerFactory
from transformers import pipeline
from docx import Document
import os
import torch

class IndigenousTranslator:
    def __init__(self, json_path):
        self.dictionary = self.load_json(json_path)
        factory = StemmerFactory()
        self.id_stemmer = factory.create_stemmer()
        self.en_stemmer = SnowballStemmer("english")
        
        # Set device for CUDA acceleration
        self.device = 0 if torch.cuda.is_available() else -1
        # Optimization configuration without device_map to avoid conflicts
        self.optimization_config = {
            "torch_dtype": torch.float16,
            "load_in_8bit": True,
            "max_memory": {"cuda:0": "3.5GB"}
        }
        
        # Use optimized translation pipelines for dyk translation
        self.en_id_translator = pipeline(
            "translation_en_to_id",
            model="Helsinki-NLP/opus-mt-en-id",
            tokenizer="Helsinki-NLP/opus-mt-en-id",
            device=self.device,
            **self.optimization_config
        )
        self.id_en_translator = pipeline(
            "translation_id_to_en",
            model="Helsinki-NLP/opus-mt-id-en",
            tokenizer="Helsinki-NLP/opus-mt-id-en",
            device=self.device,
            **self.optimization_config
        )
        self.confidence = 1.0
        self.translations = []
    
    def load_json(self, path):
        with open(path, "r") as f:
            return json.load(f)
    
    @staticmethod
    def detect_encoding(file_path):
        """Detect file encoding to avoid UnicodeDecodeError."""
        with open(file_path, "rb") as f:
            raw_data = f.read(1024)  # Read a small portion to detect encoding
            result = chardet.detect(raw_data)
        return result.get("encoding")
    
    @staticmethod
    def sanitize_text(text):
        """Remove control characters not compatible with XML."""
        return re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F]', '', text)
    
    def preprocess_text(self, text, lang):
        text = text.lower()
        if lang == "id":
            return self.id_stemmer.stem(text)
        elif lang == "en":
            return " ".join([self.en_stemmer.stem(word) for word in text.split()])
        return text
    
    def find_closest_match(self, word):
        candidates = list(self.dictionary.keys())
        if not candidates:
            return word, 0.0
        
        distances = [(c, lev_distance(word, c)) for c in candidates]
        closest = min(distances, key=lambda x: x[1])
        max_len = max(len(word), len(closest[0]))
        return closest[0], 1 - (closest[1] / max_len)
    
    def translate_word(self, word):
        if word in self.dictionary:
            return self.dictionary[word], 1.0
        else:
            closest, confidence = self.find_closest_match(word)
            if confidence > 0.6:
                return self.dictionary.get(closest, word), confidence
            return word, 0.0
    
    def translate_text(self, text, source_lang, target_lang):
        self.confidence = 1.0
        
        # For dyk translation, use the optimized translation pipelines
        if source_lang == "en" and target_lang == "dyk":
            text = self.en_id_translator(text, max_length=512)[0]["translation_text"]
        elif source_lang == "dyk" and target_lang == "en":
            # For dyk→en, first translate dyk to id using dictionary-based approach, then use pipeline for id→en
            text = self.translate_text(text, "dyk", "id")
            text = self.id_en_translator(text, max_length=512)[0]["translation_text"]
            return text
        
        # For core translation (en↔id), perform efficient word-by-word dictionary translation
        words = text.split()
        translated = []
        total_words = len(words)
        matched_words = 0
        
        for word in words:
            trans, conf = self.translate_word(word)
            translated.append(trans)
            self.confidence *= conf
            if conf > 0:
                matched_words += 1
        
        self.translations.append((text, " ".join(translated), self.confidence, total_words, len(translated)))
        self.match_rate = matched_words / total_words if total_words > 0 else 0
        return " ".join(translated)
    
    def create_report(self, filename):
        doc = Document()
        doc.add_heading("Translation Report", 0)
        
        avg_confidence = sum(conf for _, _, conf, _, _ in self.translations) / len(self.translations)
        
        header = doc.sections[0].header
        header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_paragraph.text = f"Performance Score: {avg_confidence:.1%} | Translation Rate: {self.match_rate:.1%}"
        
        for original, translated, confidence, original_word_count, target_word_count in self.translations:
            # Sanitize text to remove any XML-incompatible characters
            original = self.sanitize_text(original)
            translated = self.sanitize_text(translated)
            doc.add_paragraph(f"Original: {original}")
            doc.add_paragraph(f"Translated: {translated}")
            doc.add_paragraph(f"Confidence: {confidence:.1%}")
            doc.add_paragraph("\n")
        
        footer = doc.sections[0].footer
        footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_paragraph.text = f"Translation Origin: {original_word_count} words | Target: {target_word_count} words"
        
        doc.add_page_break()
        doc.save(filename)

if __name__ == "__main__":
    translator = IndigenousTranslator("dictionary_alt.json")
    
    while True:
        choice = input("\nEnter 'file' to load from file, 'text' for manual input, or 'exit' to quit: ").strip().lower()
        
        if choice == "exit":
            print("Exiting translator... Goodbye!")
            break
        
        if choice == "file":
            file_path = input("Enter file path: ").strip()
            if not os.path.exists(file_path):
                print("File not found. Try again.")
                continue
            
            # If the file is a DOCX, use python-docx to extract text
            if file_path.lower().endswith(".docx"):
                try:
                    doc = Document(file_path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                except Exception as e:
                    print(f"Error reading DOCX file: {e}")
                    continue
            else:
                encoding = IndigenousTranslator.detect_encoding(file_path)
                if not encoding:
                    print("Encoding could not be detected. Trying fallback encoding 'latin-1'...")
                    encoding = "latin-1"
                try:
                    with open(file_path, "r", encoding=encoding) as file:
                        text = file.read()
                except UnicodeDecodeError:
                    print(f"Error decoding file with {encoding}. Trying fallback encoding 'latin-1'...")
                    with open(file_path, "r", encoding="latin-1") as file:
                        text = file.read()
        else:
            text = input("Enter text to translate: ").strip()
        
        source_lang = input("Enter source language (id/en/dyk): ").strip().lower()
        target_lang = input("Enter target language (id/en/dyk): ").strip().lower()
        
        translation = translator.translate_text(text, source_lang, target_lang)
        translator.create_report("translation_report.docx")
        
        print(f"\nTranslation: {translation}")
        print(f"Confidence: {translator.confidence:.1%}")
