#!/usr/bin/env python3
"""
Production‐Ready ILTE–ATI Translation Engine (id ↔ dyk)
Configurable Version with Iterative Translation, Dynamic Candidate Generation, and Improved Report Output
• Loads engine settings from a JSON configuration file (if available) so that parameters aren’t hardcoded.
• Uses a broader, configurable context window for improved semantic attention.
• Relies solely on the external alt dictionary (ID → DYK) for id-dyk translation.
• For tokens not found in the alt dictionary, the engine dynamically generates candidate translations using an external
  Helsinki‑Opus‑MT en→id translator (to handle possible English words) and context‐based semantic candidate generation.
• Hierarchical normalization is performed dynamically (without hardcoded mappings) using lowercasing, regex rules, and stemming.
• Iterative translation chains processing until the output stabilizes.
• The DOCX report now includes detailed statistics and the final translated text with preserved paragraph breaks.
• Device configurations are passed to the translation pipeline.
"""

import os
import re
import json
import time
import gc
import hashlib
import logging
import psutil
import unicodedata
import numpy as np
import signal
import sys
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor, as_completed, TimeoutError
from itertools import islice
from pathlib import Path
from threading import Event, Lock

import chardet
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from Levenshtein import distance as lev_distance

# External model libraries
from Sastrawi.Stemmer.StemmerFactory import StemmerFactory
from sentence_transformers import SentenceTransformer
from transformers import MBartForConditionalGeneration, MBart50TokenizerFast, pipeline
import malaya

try:
    import torch
except ImportError:
    torch = None

# ----------------------------
# Logging configuration
# ----------------------------
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s | %(levelname)s | %(name)s | %(message)s'
)
logger = logging.getLogger("ILTE_ATI_Configurable")

# ----------------------------
# Configuration Loading
# ----------------------------
def load_config(config_path: str = "engine_config.json") -> dict:
    default_config = {
        "window_size": 10,
        "timeout": 30,
        "cache_max_age": 3600,
        "batch_size": 32,
        "thread_count": min(os.cpu_count() or 4, 8),
        "process_count": max(min((os.cpu_count() or 4) - 1, 4), 2)
    }
    if os.path.exists(config_path):
        try:
            with open(config_path, "r", encoding="utf-8") as f:
                config = json.load(f)
            logger.info("Configuration loaded from %s", config_path)
            default_config.update(config)
        except Exception as e:
            logger.error("Error loading configuration, using defaults.", exc_info=True)
    else:
        logger.info("Configuration file not found, using default settings.")
    return default_config

CONFIG = load_config()

# ----------------------------
# Resource and Cache Classes
# ----------------------------
@dataclass
class ResourceLimits:
    max_ram_gb: float = 20.0
    max_vram_gb: float = 3.8
    max_cache_gb: float = 50.0
    batch_size: int = CONFIG.get("batch_size", 32)
    thread_count: int = CONFIG.get("thread_count", min(os.cpu_count() or 4, 8))
    process_count: int = CONFIG.get("process_count", max(min((os.cpu_count() or 4) - 1, 4), 2))
    chunk_size: int = 2000
    cache_max_age: float = CONFIG.get("cache_max_age", 3600)

class CacheManager:
    """
    Caches translation results with keys including token, context, and formatting.
    """
    def __init__(self, cache_dir: Path, max_cache_gb: float, cache_filename: str = "translation_cache.json"):
        self.cache_dir = cache_dir
        self.cache_dir.mkdir(exist_ok=True)
        self.cache_file = self.cache_dir / cache_filename
        self.max_cache_gb = max_cache_gb
        self.cache: Dict[str, Dict[str, Any]] = {}
        self.lock = Lock()
        self._load_cache()

    def _load_cache(self) -> None:
        if self.cache_file.exists():
            try:
                with self.cache_file.open('r', encoding='utf-8') as f:
                    self.cache = json.load(f)
                logger.info("Cache loaded successfully.")
            except Exception as e:
                logger.error("Failed to load cache; starting with an empty cache.", exc_info=True)
                self.cache = {}
        else:
            self.cache = {}

    def _save_cache(self) -> None:
        try:
            with self.cache_file.open('w', encoding='utf-8') as f:
                json.dump(self.cache, f, ensure_ascii=False, indent=2)
            logger.info("Cache saved successfully.")
        except Exception as e:
            logger.error("Cache saving error.", exc_info=True)

    def get(self, key: str) -> Optional[Dict[str, Any]]:
        with self.lock:
            entry = self.cache.get(key)
            if entry and (time.time() - entry.get("timestamp", 0) < ResourceLimits().cache_max_age):
                return entry
            elif key in self.cache:
                del self.cache[key]
            return None

    def set(self, key: str, value: Dict[str, Any]) -> None:
        with self.lock:
            value["timestamp"] = time.time()
            self.cache[key] = value

    def clean_cache(self) -> None:
        total_size = sum(f.stat().st_size for f in self.cache_dir.glob('**/*') if f.is_file())
        while (total_size / (1024**3)) > self.max_cache_gb and self.cache:
            oldest_key = min(self.cache.items(), key=lambda x: x[1].get("timestamp", time.time()))[0]
            del self.cache[oldest_key]
            total_size = sum(f.stat().st_size for f in self.cache_dir.glob('**/*') if f.is_file())
        self._save_cache()

    def save(self) -> None:
        with self.lock:
            self._save_cache()

class ResourceManager:
    """Manages system resources and cache."""
    def __init__(self, limits: ResourceLimits):
        self.limits = limits
        self.cache_manager = CacheManager(Path("translation_cache"), limits.max_cache_gb)

    def check_memory(self) -> bool:
        memory = psutil.virtual_memory()
        used_ram_gb = memory.used / (1024**3)
        logger.debug(f"Memory used: {used_ram_gb:.2f} GB; Limit: {self.limits.max_ram_gb} GB")
        return used_ram_gb < self.limits.max_ram_gb

    def check_cache_size(self) -> bool:
        total_size = sum(f.stat().st_size for f in self.cache_manager.cache_dir.glob('**/*') if f.is_file())
        cache_size_gb = total_size / (1024**3)
        logger.debug(f"Cache size: {cache_size_gb:.2f} GB; Limit: {self.limits.max_cache_gb} GB")
        return cache_size_gb < self.limits.max_cache_gb

    def clean_cache(self) -> None:
        self.cache_manager.clean_cache()

    def release_gpu_memory(self) -> None:
        if torch is not None and torch.cuda.is_available():
            torch.cuda.empty_cache()
            logger.info("Emptied GPU cache.")

# ----------------------------
# Batching and Parallel Processing
# ----------------------------
class ContentBatcher:
    """Splits text into batches while preserving newlines and indentations."""
    def __init__(self, resource_limits):
        self.resource_limits = resource_limits

    def create_batches(self, content: str) -> List[str]:
        return content.splitlines(keepends=True)

class ParallelProcessor:
    """Handles multi-threaded and multi-process batch processing."""
    def __init__(self, resource_limits: ResourceLimits):
        self.limits = resource_limits
        self.thread_executor = ThreadPoolExecutor(max_workers=resource_limits.thread_count)
        self.process_executor = ProcessPoolExecutor(max_workers=resource_limits.process_count)

    def chunk_data(self, data: List[Any], chunk_size: Optional[int] = None) -> List[List[Any]]:
        chunk_size = chunk_size or self.limits.chunk_size
        return [list(islice(data, i, i + chunk_size)) for i in range(0, len(data), chunk_size)]

    def process_in_parallel(self, func, data: List[Any], use_processes: bool = False, timeout: int = 60) -> List[Any]:
        chunks = self.chunk_data(data)
        executor = self.process_executor if use_processes else self.thread_executor
        futures = [executor.submit(func, chunk) for chunk in chunks]
        results = []
        for future in as_completed(futures, timeout=timeout):
            try:
                results.extend(future.result(timeout=timeout))
            except Exception as e:
                logger.error("Parallel processing error.", exc_info=True)
        return results

    def shutdown(self) -> None:
        self.thread_executor.shutdown(wait=True)
        self.process_executor.shutdown(wait=True)
        logger.info("Parallel processors shutdown.")

# ----------------------------
# Scoring and Text Processing
# ----------------------------
class WeightedScorer:
    """Computes confidence scores using method weights and context factors."""
    def __init__(self):
        self.method_weights = {
            'dictionary': 1.0,
            'fst': 0.85,
            'rbmt': 0.9,
            'semantic': 0.7,
            'zero_shot': 0.5,
            'levenshtein': 0.65,
            'candidate_selection': 1.0
        }
        self.context_weights = {
            'context_window': 0.2
        }

    def calculate_confidence(self, method: str, base_score: float, context: Dict[str, float]) -> float:
        method_weight = self.method_weights.get(method, 0.5)
        context_score = sum(self.context_weights.get(k, 0) * v for k, v in context.items())
        final_score = (method_weight * base_score) + context_score
        return min(max(final_score, 0.0), 1.0)

class EnhancedTextProcessor:
    """Normalizes text, splits into sentences, and preserves formatting."""
    def __init__(self, window_size: int = CONFIG.get("window_size", 10)):
        self.window_size = window_size
        self.punctuation_pattern = re.compile(r'([.!?])')
        self.whitespace_pattern = re.compile(r'\s+')

    def normalize_text(self, text: str) -> str:
        return unicodedata.normalize('NFKC', text)

    def split_sentences(self, text: str) -> List[str]:
        normalized = self.normalize_text(text)
        lines = normalized.splitlines()
        sentences = []
        for line in lines:
            line = line.strip()
            if line:
                sentences.extend(re.split(r'(?<=[.!?])\s+', line))
        return sentences

    def get_token_context(self, tokens: List[str], idx: int) -> Dict[str, float]:
        start = max(0, idx - self.window_size)
        end = min(len(tokens), idx + self.window_size + 1)
        context_tokens = tokens[start:idx] + tokens[idx+1:end]
        return {'context_window': len(context_tokens) / (2 * self.window_size)}

# ----------------------------
# Hierarchical Normalization (Dynamic, No Hardcoded Mappings)
# ----------------------------
def hierarchical_normalize(token: str) -> str:
    """
    Dynamically normalizes a token using a hierarchy of steps:
    1. Lowercase and trim whitespace.
    2. Remove repeated characters.
    3. Apply Indonesian stemming.
    """
    token = token.lower().strip()
    token = re.sub(r'(.)\1+', r'\1', token)
    try:
        factory = StemmerFactory()
        stemmer = factory.create_stemmer()
        token = stemmer.stem(token)
    except Exception:
        pass
    return token

# ----------------------------
# Cache Key Utility
# ----------------------------
def generate_cache_key(token: str, context: Dict[str, float], formatting: str = "") -> str:
    context_str = json.dumps(context, sort_keys=True)
    key_raw = f"{token}_{context_str}_{formatting}"
    return hashlib.md5(key_raw.encode('utf-8')).hexdigest()

# ----------------------------
# External Candidate Generation for English
# ----------------------------
en_to_id_translator = pipeline(
    "translation_en_to_id",
    model="Helsinki-NLP/opus-mt-en-id",
    device=0 if torch and torch.cuda.is_available() else -1
)

def candidate_from_english(token: str) -> Optional[str]:
    if token.isalpha() and len(token) <= 5:
        try:
            result = en_to_id_translator(token)[0]["translation_text"]
            return result.lower().strip()
        except Exception as e:
            logger.warning(f"English candidate generation failed for token {token}: {e}")
    return None

# ----------------------------
# Translation Result Data Class
# ----------------------------
@dataclass
class TranslationResult:
    original: str
    translated: str
    confidence: float
    method: str
    processing_time: float

# ----------------------------
# Core Engine Implementation
# ----------------------------
class EnhancedILTETranslationEngine:
    """
    Advanced ILTE–ATI engine (id ↔ dyk) with refined attention, hierarchical normalization, and iterative translation.
    • Input is assumed to be in Indonesian (id) or DYK.
    • For id→dyk translation, if tokens are not found in the alt dictionary, dynamic candidate generation is used:
         - An external Helsinki MT en→id translator converts possible English words into standard Indonesian.
         - Context-based semantic candidate generation is also applied.
    • Iterative translation is applied until the output stabilizes.
    • Detailed statistics are recorded and output in a DOCX report, which now also includes the final translated text.
    """
    def __init__(self, dict_alt_path: str, source_lang: str = 'id', target_lang: str = 'dyk',
                 resource_limits: Optional[ResourceLimits] = None, lazy_load: bool = False):
        if source_lang not in ['id', 'dyk'] or target_lang not in ['id', 'dyk']:
            raise ValueError("Only 'id' and 'dyk' languages are supported.")
        self.shutdown_event = Event()
        self.timeout = CONFIG.get("timeout", 30)
        self.source_lang = source_lang
        self.target_lang = target_lang
        self.resource_limits = resource_limits or ResourceLimits()
        self.resource_manager = ResourceManager(self.resource_limits)
        self.parallel_processor = ParallelProcessor(self.resource_limits)
        self.text_processor = EnhancedTextProcessor()
        self.scorer = WeightedScorer()
        self.content_batcher = ContentBatcher(self.resource_limits)
        # Load alt dictionary (ID → DYK) from external JSON.
        try:
            with open(dict_alt_path, 'r', encoding='utf-8') as f:
                self.alt_dictionary = json.load(f)
            self.alt_dictionary = {str(k).strip().lower(): str(v).strip() for k, v in self.alt_dictionary.items()}
            logger.info("Alt dictionary loaded.")
        except Exception as e:
            logger.error("Error loading alt dictionary.", exc_info=True)
            raise
        self.lazy_load = lazy_load
        self._stemmer = None
        self._semantic_model = None
        self._zero_shot_model = None
        self._zero_shot_tokenizer = None
        self.semantic_embeddings: Dict[str, np.ndarray] = {}
        # Statistics tracking
        self.method_usage: Dict[str, int] = {}
        self.total_tokens: int = 0
        if not self.lazy_load:
            self._safe_initialize_models()

    def _safe_initialize_models(self) -> None:
        try:
            with ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(self._initialize_models)
                future.result(timeout=self.timeout)
        except TimeoutError:
            logger.error("Model initialization timed out")
            raise
        except Exception as e:
            logger.error("Model initialization failed", exc_info=True)
            raise

    def _initialize_models(self) -> None:
        self._load_stemmer()
        self._load_semantic_model()
        self._compute_semantic_embeddings()
        self._load_zero_shot_model()

    def _load_stemmer(self) -> None:
        try:
            factory = StemmerFactory()
            self._stemmer = factory.create_stemmer()
            logger.info("Stemmer loaded.")
        except Exception as e:
            logger.error("Failed to load stemmer.", exc_info=True)
            raise

    def _load_semantic_model(self) -> None:
        try:
            device = 'cuda' if (self.resource_limits.max_vram_gb > 0 and torch is not None and torch.cuda.is_available()) else 'cpu'
            self._semantic_model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2', device=device)
            logger.info(f"Semantic model loaded on {device}.")
        except Exception as e:
            logger.error("Failed to load semantic model.", exc_info=True)
            raise

    def _compute_semantic_embeddings(self) -> None:
        try:
            words = list(self.alt_dictionary.keys())
            chunks = self.parallel_processor.chunk_data(words)
            embeddings: Dict[str, np.ndarray] = {}
            futures = []
            for chunk in chunks:
                future = self.parallel_processor.thread_executor.submit(
                    self._semantic_model.encode,
                    chunk,
                    batch_size=self.resource_limits.batch_size,
                    show_progress_bar=False
                )
                futures.append((chunk, future))
            for chunk, future in futures:
                try:
                    batch_embeddings = future.result(timeout=120)
                    for word, emb in zip(chunk, batch_embeddings):
                        embeddings[word] = emb
                except Exception as e:
                    logger.error("Error computing embeddings for a chunk.", exc_info=True)
            self.semantic_embeddings = embeddings
            logger.info("Semantic embeddings computed.")
        except Exception as e:
            logger.error("Failed to compute semantic embeddings.", exc_info=True)

    def _load_zero_shot_model(self) -> None:
        try:
            self._zero_shot_model = MBartForConditionalGeneration.from_pretrained(
                'facebook/mbart-large-50-many-to-many-mmt',
                low_cpu_mem_usage=True
            )
            self._zero_shot_tokenizer = MBart50TokenizerFast.from_pretrained(
                'facebook/mbart-large-50-many-to-many-mmt'
            )
            logger.info("Zero-shot model and tokenizer loaded.")
        except Exception as e:
            logger.error("Failed to load zero-shot model.", exc_info=True)
            raise

    @property
    def stemmer(self):
        if self._stemmer is None:
            self._load_stemmer()
        return self._stemmer

    @property
    def semantic_model(self):
        if self._semantic_model is None:
            self._load_semantic_model()
        return self._semantic_model

    @property
    def zero_shot_model(self):
        if self._zero_shot_model is None:
            self._load_zero_shot_model()
        return self._zero_shot_model

    @property
    def zero_shot_tokenizer(self):
        if self._zero_shot_tokenizer is None:
            self._load_zero_shot_model()
        return self._zero_shot_tokenizer

    def _translate_core(self, core: str, context: Dict[str, float], formatting: str, start_time: float) -> TranslationResult:
        cache_key = generate_cache_key(core, context, formatting)
        cached = self.resource_manager.cache_manager.get(cache_key)
        if cached:
            return TranslationResult(core, cached["translated"], cached["confidence"], cached["method"], time.time() - start_time)
        self.total_tokens += 1
        norm_core = hierarchical_normalize(core)
        if norm_core in self.alt_dictionary:
            conf = self.scorer.calculate_confidence('dictionary', 1.0, context)
            result = TranslationResult(norm_core, self.alt_dictionary[norm_core], conf, 'dictionary', time.time() - start_time)
            self.method_usage['dictionary'] = self.method_usage.get('dictionary', 0) + 1
            self.resource_manager.cache_manager.set(cache_key, {
                "translated": result.translated,
                "confidence": result.confidence,
                "method": result.method
            })
            return result
        
        # Dynamic candidate generation for possible English words
        en_candidate = candidate_from_english(norm_core)
        if en_candidate and en_candidate in self.alt_dictionary:
            conf = self.scorer.calculate_confidence('candidate_selection', 1.0, context)
            result = TranslationResult(norm_core, self.alt_dictionary[en_candidate], conf, 'candidate_selection', time.time() - start_time)
            self.method_usage['candidate_selection'] = self.method_usage.get('candidate_selection', 0) + 1
            self.resource_manager.cache_manager.set(cache_key, {
                "translated": result.translated,
                "confidence": result.confidence,
                "method": result.method
            })
            return result
        
        # Fallback: use complex translation module
        try:
            with ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(self._complex_translation, norm_core, norm_core, context, formatting, start_time)
                result = future.result(timeout=self.timeout)
                self.resource_manager.cache_manager.set(cache_key, {
                    "translated": result.translated,
                    "confidence": result.confidence,
                    "method": result.method
                })
                self.method_usage[result.method] = self.method_usage.get(result.method, 0) + 1
                return result
        except TimeoutError:
            logger.warning(f"Translation timed out for token: {norm_core}")
            self.method_usage['timeout_fallback'] = self.method_usage.get('timeout_fallback', 0) + 1
            return TranslationResult(norm_core, norm_core, 0.1, 'timeout_fallback', time.time() - start_time)
        except Exception as e:
            logger.error(f"Translation failed for token: {norm_core}", exc_info=True)
            self.method_usage['error_fallback'] = self.method_usage.get('error_fallback', 0) + 1
            return TranslationResult(norm_core, norm_core, 0.1, 'error_fallback', time.time() - start_time)

    def translate_token(self, token: str, tokens: List[str], idx: int, formatting: str = "") -> TranslationResult:
        if re.fullmatch(r'\W+', token):
            return TranslationResult(token, token, 1.0, 'punctuation', 0.0)
        prefix_match = re.match(r'^(\s+)', token)
        suffix_match = re.search(r'(\s+)$', token)
        pre = prefix_match.group(1) if prefix_match else ""
        suf = suffix_match.group(1) if suffix_match else ""
        core = token[len(pre):len(token)-len(suf)] if suf else token[len(pre):]
        core_clean = core.strip().lower()
        context = self.text_processor.get_token_context(tokens, idx)
        result = self._translate_core(core_clean, context, formatting, start_time=time.time())
        translated = pre + result.translated + suf
        return TranslationResult(token, translated, result.confidence, result.method, result.processing_time)

    def _complex_translation(self, word: str, processed: str, context: Dict[str, float], formatting: str, start_time: float) -> TranslationResult:
        try:
            fst_candidate, fst_conf = self.apply_fst_rules(processed)
            if fst_candidate:
                conf = self.scorer.calculate_confidence('fst', fst_conf, context)
                return TranslationResult(word, fst_candidate, conf, 'fst', time.time() - start_time)
            rbmt_candidate, rbmt_conf = self.apply_rbmt_rules(processed)
            if rbmt_candidate:
                conf = self.scorer.calculate_confidence('rbmt', rbmt_conf, context)
                return TranslationResult(word, rbmt_candidate, conf, 'rbmt', time.time() - start_time)
            try:
                word_emb = self.semantic_model.encode([processed])[0]
                candidate_scores = []
                for dict_word, emb in self.semantic_embeddings.items():
                    norm_product = np.linalg.norm(word_emb) * np.linalg.norm(emb)
                    if norm_product == 0:
                        continue
                    cos_sim = np.dot(word_emb, emb) / norm_product
                    lev_sim = 1 - (lev_distance(processed, dict_word) / max(len(processed), len(dict_word)))
                    combined = 0.6 * cos_sim + 0.4 * lev_sim
                    candidate_scores.append((dict_word, combined))
                candidate_scores.sort(key=lambda x: x[1], reverse=True)
                top_candidates = [cand for cand in candidate_scores if cand[1] > 0.7][:3]
                if top_candidates:
                    best_candidate, best_score = top_candidates[0]
                    conf = self.scorer.calculate_confidence('semantic', best_score, context)
                    return TranslationResult(word, self.alt_dictionary.get(best_candidate, best_candidate), conf, 'semantic', time.time() - start_time)
            except Exception as e:
                logger.warning(f"Semantic translation failed for '{word}': {e}", exc_info=True)
            try:
                inputs = self.zero_shot_tokenizer(processed, return_tensors="pt", padding=True)
                forced_lang = "id_ID"
                self.zero_shot_tokenizer.src_lang = "id_ID"
                generated = self.zero_shot_model.generate(
                    **inputs,
                    forced_bos_token_id=self.zero_shot_tokenizer.lang_code_to_id[forced_lang],
                    max_length=128,
                    num_beams=8,
                    length_penalty=0.6,
                    max_time=10.0
                )
                zs_trans = self.zero_shot_tokenizer.batch_decode(generated, skip_special_tokens=True)[0]
                conf = self.scorer.calculate_confidence('zero_shot', 0.5, context)
                return TranslationResult(word, zs_trans, conf, 'zero_shot', time.time() - start_time)
            except Exception as e:
                logger.warning(f"Zero-shot translation failed for '{word}': {e}", exc_info=True)
            return TranslationResult(word, word, 0.1, 'fallback', time.time() - start_time)
        finally:
            if torch is not None and torch.cuda.is_available():
                torch.cuda.empty_cache()

    def apply_fst_rules(self, word: str) -> Tuple[Optional[str], float]:
        rules = [
            (r"([a-zA-Z]+)s$", r"\1", 0.85),
            (r"([a-zA-Z]+)(an|wan)$", r"\1", 0.85),
            (r"([a-zA-Z]+)nya$", r"\1", 0.85),
            (r"^(un|in|ter)([a-zA-Z]+)$", r"\2", 0.85),
            (r"^di([a-zA-Z]+)$", r"\1", 0.85),
            (r"^ke([a-zA-Z]+)$", r"\1", 0.85),
            (r"^me([a-zA-Z]+)$", r"\1", 0.85),
            (r"^pe([a-zA-Z]+)$", r"\1", 0.85)
        ]
        for pattern, replacement, confidence in rules:
            candidate = re.sub(pattern, replacement, word, flags=re.IGNORECASE)
            if candidate != word and candidate in self.alt_dictionary:
                return candidate, confidence
        return None, 0.0

    def apply_rbmt_rules(self, word: str) -> Tuple[Optional[str], float]:
        rbmt_rules = [
            (r"\bselamat pagi\b", "selamat pagi", 0.9),
            (r"\bapa kabar\b", "apa kabar", 0.9)
        ]
        for pattern, replacement, conf in rbmt_rules:
            if re.search(pattern, word, re.IGNORECASE):
                return replacement, conf
        if word.endswith("ing"):
            candidate = word[:-3]
            if candidate in self.alt_dictionary:
                return self.alt_dictionary[candidate], 0.8
        return None, 0.0

    def translate_sentence(self, sentence: str) -> str:
        tokens = re.findall(r'\s+|\S+', sentence)
        results = [None] * len(tokens)
        with ThreadPoolExecutor(max_workers=self.resource_limits.thread_count) as executor:
            futures = {executor.submit(self.translate_token, token, tokens, idx, ""): idx for idx, token in enumerate(tokens)}
            for future in as_completed(futures, timeout=60):
                idx = futures[future]
                try:
                    res = future.result(timeout=60)
                    results[idx] = res.translated
                except Exception as e:
                    logger.error("Token translation error.", exc_info=True)
                    results[idx] = tokens[idx]
        return "".join(results)

    def translate_text(self, text: str, source_lang: str, target_lang: str) -> str:
        logger.info(f"Translating text from {source_lang} to {target_lang}.")
        lines = text.splitlines(keepends=True)
        translated_lines = []
        for line in lines:
            if line.strip():
                translated_line = self.translate_sentence(line)
            else:
                translated_line = line
            translated_lines.append(translated_line)
        overall_translation = "".join(translated_lines)
        gc.collect()
        return overall_translation

    def iterative_translate_text(self, text: str, source_lang: str, target_lang: str, max_iterations: int = 3) -> str:
        prev = text
        for i in range(max_iterations):
            new_text = self.translate_text(prev, source_lang, target_lang)
            if new_text.strip() == prev.strip():
                break
            prev = new_text
        return prev

    def create_report(self, translation_text: str, filename: str) -> None:
        doc = Document()
        doc.add_heading("ILTE–ATI Translation Report", 0)
        # Statistics header
        avg_confidence = 0.0  # Placeholder
        method_counts = self.method_usage.copy()
        total = self.total_tokens if self.total_tokens > 0 else 1
        for method, count in method_counts.items():
            method_counts[method] = (count / total) * 100
        p = doc.add_paragraph()
        p.add_run("Average Confidence Score: ").bold = True
        p.add_run(f"{avg_confidence:.2f}")
        p = doc.add_paragraph()
        p.add_run("Total Tokens Processed: ").bold = True
        p.add_run(f"{self.total_tokens}")
        p = doc.add_paragraph()
        p.add_run("Method Usage Rates (%):").bold = True
        for method, usage in method_counts.items():
            doc.add_paragraph(f"{method}: {usage:.1f}%", style='List Bullet')
        non_fallback = self.total_tokens - self.method_usage.get('fallback', 0) - \
                       self.method_usage.get('timeout_fallback', 0) - self.method_usage.get('error_fallback', 0)
        translation_rate = (non_fallback / total) * 100
        p = doc.add_paragraph()
        p.add_run("Translation Rate: ").bold = True
        p.add_run(f"{translation_rate:.1f}%")
        doc.add_page_break()
        # Append the final translated text preserving paragraph breaks
        paragraphs = translation_text.splitlines()
        for para in paragraphs:
            doc.add_paragraph(para)
        doc.save(filename)
        logger.info(f"Report saved to {filename}")

    def cleanup(self) -> None:
        self.parallel_processor.shutdown()
        self.resource_manager.cache_manager.save()
        self.resource_manager.clean_cache()
        self.resource_manager.release_gpu_memory()
        gc.collect()

# ----------------------------
# Helper: File Encoding Detection
# ----------------------------
def detect_encoding(file_path: str, read_bytes: int = 100000) -> str:
    try:
        with open(file_path, "rb") as f:
            raw_data = f.read(read_bytes)
        detected = chardet.detect(raw_data)
        encoding = detected.get("encoding", "utf-8")
        logger.info(f"Detected encoding: {encoding}")
        return encoding
    except Exception as e:
        logger.error("Encoding detection error.", exc_info=True)
        return "utf-8"

# ----------------------------
# Engine Runner: Interactive/File Modes
# ----------------------------
class ILTEEngineRunner:
    """
    Main interface for the refined ILTE–ATI engine.
    Supports interactive and file input while preserving formatting.
    The DOCX report includes detailed statistics and the final translated text.
    Iterative translation is applied so that inputs like "Hi" yield "Halo" then "denga".
    """
    def __init__(self, engine: EnhancedILTETranslationEngine):
        self.engine = engine
        signal.signal(signal.SIGINT, self._handle_interrupt)
        signal.signal(signal.SIGTERM, self._handle_interrupt)

    def _handle_interrupt(self, signum, frame):
        print("\nReceived interrupt signal. Cleaning up...")
        self.engine.shutdown_event.set()
        self.engine.cleanup()
        print("Cleanup complete. Exiting...")
        sys.exit(0)

    def process_text(self, text: str) -> List[str]:
        lines = text.splitlines()
        results = []
        for line in lines:
            result = self.engine.iterative_translate_text(line, self.engine.source_lang, self.engine.target_lang)
            results.append(result)
        return results

    def process_file(self, file_path: str) -> List[str]:
        try:
            safe_path = Path(file_path).resolve()
            if not safe_path.exists():
                logger.error("File not found.")
                return []
            encoding = detect_encoding(str(safe_path))
            if safe_path.suffix.lower() == ".docx":
                doc = Document(str(safe_path))
                text = "\n".join([para.text for para in doc.paragraphs])
            else:
                with safe_path.open('r', encoding=encoding, errors='replace') as file:
                    text = file.read()
            return self.process_text(text)
        except Exception as e:
            logger.error("File processing error.", exc_info=True)
            return []

    def generate_report(self, translation_text: str, filename: str = "translation_report.docx") -> None:
        self.engine.create_report(translation_text, filename)

    def interactive_mode(self) -> None:
        print("\n--- Interactive Translation Mode (id ↔ dyk) ---")
        while not self.engine.shutdown_event.is_set():
            try:
                text = input("\nEnter text to translate (or type 'exit' to quit): ")
                if text.lower().strip() == "exit":
                    break
                print("Translating...")
                results = self.process_text(text)
                final_translation = "\n".join(results)
                print("\n--- Translation Results ---")
                print(final_translation)
                self.generate_report(final_translation)
            except KeyboardInterrupt:
                print("\nTranslation interrupted by user.")
                break
            except Exception as e:
                print(f"\nError occurred: {str(e)}")
                logger.error("Interactive mode error", exc_info=True)
                break

    def file_mode(self) -> None:
        file_path = input("\nEnter file path: ").strip()
        if not os.path.exists(file_path):
            print("File not found.")
            return
        print("\nDetecting file encoding...")
        encoding = detect_encoding(file_path)
        print(f"Detected encoding: {encoding}")
        print("Processing file...")
        results = self.process_file(file_path)
        final_translation = "\n".join(results)
        print("\n--- Preview (first 5 lines) ---")
        for line in final_translation.splitlines()[:5]:
            print(line)
        self.generate_report(final_translation)

    def main(self) -> None:
        print("\n=== Welcome to the Refined ILTE–ATI Translation Engine ===")
        while True:
            mode = input("\nChoose mode: [1] Direct Input  [2] File Input  [3] Exit: ").strip()
            if mode == "1":
                self.interactive_mode()
            elif mode == "2":
                self.file_mode()
            elif mode == "3":
                print("Exiting translation engine...")
                break
            else:
                print("Invalid choice. Try again.")

# ----------------------------
# Main Execution
# ----------------------------
if __name__ == "__main__":
    engine = None
    try:
        dict_alt_path = "dictionary_alt.json"
        # For id→dyk translation, we only need one dictionary file.
        engine = EnhancedILTETranslationEngine(dict_alt_path, source_lang='id', target_lang='dyk')
        runner = ILTEEngineRunner(engine)
        runner.main()
    except Exception as main_exc:
        logger.critical("Fatal error in ILTE–ATI engine execution.", exc_info=True)
    finally:
        if engine is not None:
            engine.cleanup()
            logger.info("ILTE–ATI engine cleanup complete.")
        sys.exit(0)
