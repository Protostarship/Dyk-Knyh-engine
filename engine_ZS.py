#!/usr/bin/env python3
"""
Refactored ILTE-ZS Translation Engine with RBMT & FST Improvements

Key improvements:
  • Resource management: Executors are shutdown via explicit cleanup; GPU memory is released if available.
  • Error handling & logging: Detailed, structured logging with timeouts and shutdown control.
  • Model loading: Heavy models are loaded via a safe initialization method with timeout protection.
  • Translation quality: Enhanced pipeline now uses dictionary lookup, FST rules, RBMT rules, Malay cognate detection,
    semantic similarity (with Levenshtein), and zero-shot translation.
  • Cache management: Uses a thread-safe CacheManager with JSON serialization and robust cache key generation.
  • Concurrency: Thread locks, timeouts, and shutdown events protect shared resources and prevent stuck progress.
  • I/O enhancements: Expanded encoding detection, Unicode normalization, and safe file handling.
  • Input sanitation: Text is normalized (including proper splitting of punctuation) and each word is cleaned of extraneous punctuation before translation. 
  • Security: JSON is used for caching (avoiding pickle risks), and input sanitization is performed.
"""

import os
import re
import json
import time
import gc
import hashlib
import logging
import psutil
import unicodedata
import numpy as np
import signal
import sys
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor, as_completed, TimeoutError
from itertools import islice
from pathlib import Path
from threading import Event

import chardet
from docx import Document
from Levenshtein import distance as lev_distance

# External model libraries
from Sastrawi.Stemmer.StemmerFactory import StemmerFactory
from sentence_transformers import SentenceTransformer
from transformers import MBartForConditionalGeneration, MBart50TokenizerFast
import malaya

# If torch is available, we will use it for GPU memory cleanup.
try:
    import torch
except ImportError:
    torch = None

# Configure structured logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s | %(levelname)s | %(name)s | %(message)s'
)
logger = logging.getLogger("ILTE_ZS")

@dataclass
class ResourceLimits:
    """Resource management configuration."""
    max_ram_gb: float = 12.0
    max_vram_gb: float = 3.5
    max_cache_gb: float = 30.0
    batch_size: int = 32
    thread_count: int = min(os.cpu_count() or 4, 8)
    process_count: int = max(min((os.cpu_count() or 4) - 1, 4), 2)
    chunk_size: int = 1000
    cache_max_age: float = 3600.0  # seconds

class CacheManager:
    """
    Manages caching of translation results using JSON serialization.
    Uses a lock for thread-safe operations.
    """
    def __init__(self, cache_dir: Path, max_cache_gb: float, cache_filename: str = "translation_cache.json"):
        self.cache_dir = cache_dir
        self.cache_dir.mkdir(exist_ok=True)
        self.cache_file = self.cache_dir / cache_filename
        self.max_cache_gb = max_cache_gb
        self.cache: Dict[str, Dict[str, Any]] = {}
        self.lock = __import__("threading").Lock()
        self._load_cache()

    def _load_cache(self) -> None:
        if self.cache_file.exists():
            try:
                with self.cache_file.open('r', encoding='utf-8') as f:
                    self.cache = json.load(f)
                logger.info("Cache loaded successfully.")
            except Exception as e:
                logger.error("Failed to load cache; starting with an empty cache.", exc_info=True)
                self.cache = {}
        else:
            self.cache = {}

    def _save_cache(self) -> None:
        try:
            with self.cache_file.open('w', encoding='utf-8') as f:
                json.dump(self.cache, f, ensure_ascii=False, indent=2)
            logger.info("Cache saved successfully.")
        except Exception as e:
            logger.error("Cache saving error.", exc_info=True)

    def get(self, key: str) -> Optional[Dict[str, Any]]:
        with self.lock:
            entry = self.cache.get(key)
            if entry and (time.time() - entry.get("timestamp", 0) < ResourceLimits().cache_max_age):
                return entry
            else:
                if key in self.cache:
                    del self.cache[key]
                return None

    def set(self, key: str, value: Dict[str, Any]) -> None:
        with self.lock:
            value["timestamp"] = time.time()
            self.cache[key] = value

    def clean_cache(self) -> None:
        total_size = sum(f.stat().st_size for f in self.cache_dir.glob('**/*') if f.is_file())
        while (total_size / (1024**3)) > self.max_cache_gb and self.cache:
            oldest_key = min(self.cache.items(), key=lambda x: x[1].get("timestamp", time.time()))[0]
            del self.cache[oldest_key]
            total_size = sum(f.stat().st_size for f in self.cache_dir.glob('**/*') if f.is_file())
        self._save_cache()

    def save(self) -> None:
        with self.lock:
            self._save_cache()

def generate_cache_key(word: str, context: Dict[str, float]) -> str:
    """
    Generate a robust cache key using md5 hash.
    """
    context_str = json.dumps(context, sort_keys=True)
    key_raw = f"{word}_{context_str}"
    return hashlib.md5(key_raw.encode('utf-8')).hexdigest()

class ResourceManager:
    """Manages system resources and constraints."""
    def __init__(self, limits: ResourceLimits):
        self.limits = limits
        self.cache_manager = CacheManager(Path("translation_cache"), limits.max_cache_gb)

    def check_memory(self) -> bool:
        memory = psutil.virtual_memory()
        used_ram_gb = memory.used / (1024**3)
        logger.debug(f"Memory used: {used_ram_gb:.2f} GB; Limit: {self.limits.max_ram_gb} GB")
        return used_ram_gb < self.limits.max_ram_gb

    def check_cache_size(self) -> bool:
        total_size = sum(f.stat().st_size for f in self.cache_manager.cache_dir.glob('**/*') if f.is_file())
        cache_size_gb = total_size / (1024**3)
        logger.debug(f"Cache size: {cache_size_gb:.2f} GB; Limit: {self.limits.max_cache_gb} GB")
        return cache_size_gb < self.limits.max_cache_gb

    def clean_cache(self) -> None:
        self.cache_manager.clean_cache()

    def release_gpu_memory(self) -> None:
        """Attempt to free GPU memory if using CUDA."""
        if torch is not None and torch.cuda.is_available():
            torch.cuda.empty_cache()
            logger.info("Emptied GPU cache.")

class ContentBatcher:
    """Handles batching of large text content."""
    def __init__(self, resource_limits: ResourceLimits):
        self.limits = resource_limits

    def estimate_optimal_batch_size(self, total_content_size: int) -> int:
        available_memory_gb = psutil.virtual_memory().available / (1024**3)
        cpu_count = self.limits.process_count
        memory_based_size = int((available_memory_gb * (1024**3)) / (cpu_count * 2))
        return min(max(memory_based_size // 100, 500), 5000)

    def create_batches(self, content: str) -> List[str]:
        sentences = re.split(r'(?<=[.!?])\s+', content)
        return sentences

class ParallelProcessor:
    """Handles parallel processing using threads or processes."""
    def __init__(self, resource_limits: ResourceLimits):
        self.limits = resource_limits
        self.thread_executor = ThreadPoolExecutor(max_workers=resource_limits.thread_count)
        self.process_executor = ProcessPoolExecutor(max_workers=resource_limits.process_count)

    def chunk_data(self, data: List[Any], chunk_size: Optional[int] = None) -> List[List[Any]]:
        chunk_size = chunk_size or self.limits.chunk_size
        return [list(islice(data, i, i + chunk_size)) for i in range(0, len(data), chunk_size)]

    def process_in_parallel(self, func, data: List[Any], use_processes: bool = False, timeout: int = 60) -> List[Any]:
        chunks = self.chunk_data(data)
        executor = self.process_executor if use_processes else self.thread_executor
        futures = [executor.submit(func, chunk) for chunk in chunks]
        results = []
        for future in as_completed(futures, timeout=timeout):
            try:
                results.extend(future.result(timeout=timeout))
            except Exception as e:
                logger.error("Parallel processing error.", exc_info=True)
        return results

    def shutdown(self) -> None:
        """Gracefully shutdown both thread and process executors."""
        self.thread_executor.shutdown(wait=True)
        self.process_executor.shutdown(wait=True)
        logger.info("Parallel processors shutdown.")

class WeightedScorer:
    """Calculates weighted confidence scores based on method and context."""
    def __init__(self):
        self.method_weights = {
            'dictionary': 1.0,
            'fst': 0.85,
            'rbmt': 0.9,
            'malay_cognate': 0.75,
            'semantic': 0.7,
            'zero_shot': 0.5,
            'levenshtein': 0.65
        }
        self.context_weights = {
            'sentence_position': 0.1,
            'surrounding_words': 0.15,
            'grammar_structure': 0.2
        }

    def calculate_confidence(self, method: str, base_score: float, context: Dict[str, float]) -> float:
        method_weight = self.method_weights.get(method, 0.5)
        context_score = sum(self.context_weights.get(k, 0) * v for k, v in context.items())
        final_score = (method_weight * base_score) + context_score
        return min(max(final_score, 0.0), 1.0)

class EnhancedTextProcessor:
    """Provides text normalization, Unicode sanitization, and sentence splitting."""
    def __init__(self, parallel_processor: ParallelProcessor):
        self.parallel_processor = parallel_processor
        self.punctuation_pattern = re.compile(r'([.!?])')
        self.whitespace_pattern = re.compile(r'\s+')

    def normalize_text(self, text: str) -> str:
        text = unicodedata.normalize('NFKC', text)
        # Insert spaces around punctuation so they can be separated later.
        text = self.punctuation_pattern.sub(r' \1 ', text)
        text = self.whitespace_pattern.sub(' ', text)
        return text.strip().lower()

    def split_sentences(self, text: str) -> List[str]:
        # First normalize the entire text.
        normalized = self.normalize_text(text)
        # Then split on punctuation that ends a sentence.
        sentences = re.split(r'(?<=[.!?])\s+', normalized)
        return sentences

@dataclass
class TranslationResult:
    """Translation result structure."""
    original: str
    translated: str
    confidence: float
    method: str
    processing_time: float

class EnhancedILTEZSTranslationEngine:
    """
    Core translation engine with multiple layers, lazy model loading,
    resource-optimized processing, and a robust pipeline (including RBMT and FST rules).
    """
    def __init__(self, dictionary_path: str, source_lang: str = 'id', target_lang: str = 'dyk',
                 resource_limits: Optional[ResourceLimits] = None, lazy_load: bool = False):
        self.shutdown_event = Event()
        self.timeout = 30  # seconds timeout for operations
        self.source_lang = source_lang
        self.target_lang = target_lang
        self.resource_limits = resource_limits or ResourceLimits()
        self.resource_manager = ResourceManager(self.resource_limits)
        self.parallel_processor = ParallelProcessor(self.resource_limits)
        self.text_processor = EnhancedTextProcessor(self.parallel_processor)
        self.scorer = WeightedScorer()
        self.content_batcher = ContentBatcher(self.resource_limits)
        self.dictionary = self._load_dictionary(dictionary_path)
        self.lazy_load = lazy_load
        # Placeholders for heavy models.
        self._stemmer = None
        self._semantic_model = None
        self._zero_shot_model = None
        self._zero_shot_tokenizer = None
        self._malay_model = None
        self.semantic_embeddings: Dict[str, np.ndarray] = {}
        if not self.lazy_load:
            self._safe_initialize_models()

    def _load_dictionary(self, path: str) -> Dict[str, str]:
        try:
            with open(path, 'r', encoding='utf-8') as f:
                dictionary = json.load(f)
            validated = {str(k).strip().lower(): str(v).strip() for k, v in dictionary.items() if isinstance(k, str) and isinstance(v, str)}
            if self.source_lang == 'dyk' and self.target_lang == 'id':
                validated = {v: k for k, v in validated.items()}
            return validated
        except Exception as e:
            logger.error("Dictionary loading error.", exc_info=True)
            raise

    def _safe_initialize_models(self) -> None:
        """Initialize models with timeout protection."""
        try:
            with ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(self._initialize_models)
                future.result(timeout=self.timeout)
        except TimeoutError:
            logger.error("Model initialization timed out")
            raise
        except Exception as e:
            logger.error("Model initialization failed", exc_info=True)
            raise

    def _initialize_models(self) -> None:
        try:
            self._load_stemmer()
            self._load_semantic_model()
            self._compute_semantic_embeddings()
            self._load_zero_shot_model()
            self._load_malay_model()
        except Exception as e:
            logger.error("Model initialization error.", exc_info=True)
            raise

    def _load_stemmer(self) -> None:
        try:
            factory = StemmerFactory()
            self._stemmer = factory.create_stemmer()
            logger.info("Stemmer loaded.")
        except Exception as e:
            logger.error("Failed to load stemmer.", exc_info=True)
            raise

    def _load_semantic_model(self) -> None:
        try:
            device = 'cuda' if (self.resource_limits.max_vram_gb > 0 and torch is not None and torch.cuda.is_available()) else 'cpu'
            self._semantic_model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2', device=device)
            logger.info(f"Semantic model loaded on {device}.")
        except Exception as e:
            logger.error("Failed to load semantic model.", exc_info=True)
            raise

    def _compute_semantic_embeddings(self) -> None:
        try:
            words = list(self.dictionary.keys())
            chunks = self.parallel_processor.chunk_data(words)
            embeddings: Dict[str, np.ndarray] = {}
            futures = []
            for chunk in chunks:
                future = self.parallel_processor.thread_executor.submit(
                    self._semantic_model.encode,
                    chunk,
                    batch_size=self.resource_limits.batch_size,
                    show_progress_bar=False
                )
                futures.append((chunk, future))
            for chunk, future in futures:
                try:
                    batch_embeddings = future.result(timeout=120)
                    for word, emb in zip(chunk, batch_embeddings):
                        embeddings[word] = emb
                except Exception as e:
                    logger.error("Error computing embeddings for a chunk.", exc_info=True)
            self.semantic_embeddings = embeddings
            logger.info("Semantic embeddings computed.")
        except Exception as e:
            logger.error("Failed to compute semantic embeddings.", exc_info=True)

    def _load_zero_shot_model(self) -> None:
        try:
            self._zero_shot_model = MBartForConditionalGeneration.from_pretrained(
                'facebook/mbart-large-50-many-to-many-mmt',
                low_cpu_mem_usage=True
            )
            self._zero_shot_tokenizer = MBart50TokenizerFast.from_pretrained(
                'facebook/mbart-large-50-many-to-many-mmt'
            )
            logger.info("Zero-shot model and tokenizer loaded.")
        except Exception as e:
            logger.error("Failed to load zero-shot model.", exc_info=True)
            raise

    def _load_malay_model(self) -> None:
        try:
            self._malay_model = malaya.language_detection.fasttext()
            logger.info("Malay language detection model loaded.")
        except Exception as e:
            logger.error("Failed to load Malay model.", exc_info=True)

    @property
    def stemmer(self):
        if self._stemmer is None:
            self._load_stemmer()
        return self._stemmer

    @property
    def semantic_model(self):
        if self._semantic_model is None:
            self._load_semantic_model()
        return self._semantic_model

    @property
    def zero_shot_model(self):
        if self._zero_shot_model is None:
            self._load_zero_shot_model()
        return self._zero_shot_model

    @property
    def zero_shot_tokenizer(self):
        if self._zero_shot_tokenizer is None:
            self._load_zero_shot_model()
        return self._zero_shot_tokenizer

    @property
    def malay_model(self):
        if self._malay_model is None:
            self._load_malay_model()
        return self._malay_model

    def _translate_core(self, core: str, context: Optional[Dict[str, float]], start_time: float) -> TranslationResult:
        context = context or {}
        cache_key = generate_cache_key(core, context)
        cached = self.resource_manager.cache_manager.get(cache_key)
        if cached:
            return TranslationResult(core, cached["translated"], cached["confidence"], cached["method"], 0.0)
        if core in self.dictionary:
            conf = self.scorer.calculate_confidence('dictionary', 1.0, context)
            result = TranslationResult(core, self.dictionary[core], conf, 'dictionary', time.time() - start_time)
            self.resource_manager.cache_manager.set(cache_key, {
                "translated": result.translated,
                "confidence": result.confidence,
                "method": result.method
            })
            return result
        try:
            with ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(self._complex_translation, core, core, context, start_time)
                result = future.result(timeout=self.timeout)
                self.resource_manager.cache_manager.set(cache_key, {
                    "translated": result.translated,
                    "confidence": result.confidence,
                    "method": result.method
                })
                return result
        except TimeoutError:
            logger.warning(f"Translation timed out for word: {core}")
            return TranslationResult(core, core, 0.1, 'timeout_fallback', time.time() - start_time)
        except Exception as e:
            logger.error(f"Translation failed for word: {core}", exc_info=True)
            return TranslationResult(core, core, 0.1, 'error_fallback', time.time() - start_time)

    def translate_word(self, word: str, context: Optional[Dict[str, float]] = None) -> TranslationResult:
        """
        Translate a word with proper cleaning: remove any leading/trailing punctuation,
        process the core token, and then reattach punctuation.
        """
        if re.fullmatch(r'\W+', word):
            return TranslationResult(word, word, 1.0, 'punctuation', 0.0)
        prefix_match = re.match(r'^(\W+)', word)
        suffix_match = re.search(r'(\W+)$', word)
        pre = prefix_match.group(1) if prefix_match else ""
        suf = suffix_match.group(1) if suffix_match else ""
        core = word[len(pre):len(word)-len(suf)] if suf else word[len(pre):]
        # Normalize and clean the core token.
        core_clean = self.text_processor.normalize_text(core).strip()
        result = self._translate_core(core_clean, context, start_time=time.time())
        translated = pre + result.translated + suf
        return TranslationResult(word, translated, result.confidence, result.method, result.processing_time)

    def _complex_translation(self, word: str, processed: str, context: Dict[str, float], start_time: float) -> TranslationResult:
        """Handle complex translation methods with proper resource cleanup."""
        try:
            # 1. Apply FST rules
            fst_candidate, fst_conf = self.apply_fst_rules(processed)
            if fst_candidate:
                conf = self.scorer.calculate_confidence('fst', fst_conf, context)
                return TranslationResult(word, fst_candidate, conf, 'fst', time.time() - start_time)
            # 2. Apply RBMT rules
            rbmt_candidate, rbmt_conf = self.apply_rbmt_rules(processed)
            if rbmt_candidate:
                conf = self.scorer.calculate_confidence('rbmt', rbmt_conf, context)
                return TranslationResult(word, rbmt_candidate, conf, 'rbmt', time.time() - start_time)
            # 3. Malay cognate detection
            try:
                if self.malay_model is not None:
                    malay_pred = self.malay_model.predict_proba([processed])
                    if "malay" in malay_pred[0] and malay_pred[0]["malay"] > 0.8:
                        stemmed = self.stemmer.stem(processed)
                        if stemmed in self.dictionary:
                            conf = self.scorer.calculate_confidence('malay_cognate', malay_pred[0]["malay"], context)
                            return TranslationResult(word, self.dictionary[stemmed], conf, 'malay_cognate', time.time() - start_time)
            except Exception as e:
                logger.warning(f"Malay cognate detection failed for '{word}': {e}", exc_info=True)
            # 4. Semantic similarity with Levenshtein fallback
            try:
                word_emb = self.semantic_model.encode([processed])[0]
                best_score = -1.0
                best_candidate = None
                for dict_word, emb in self.semantic_embeddings.items():
                    norm_product = np.linalg.norm(word_emb) * np.linalg.norm(emb)
                    if norm_product == 0:
                        continue
                    cos_sim = np.dot(word_emb, emb) / norm_product
                    lev_sim = 1 - (lev_distance(processed, dict_word) / max(len(processed), len(dict_word)))
                    combined = 0.6 * cos_sim + 0.4 * lev_sim
                    if combined > best_score:
                        best_score = combined
                        best_candidate = dict_word
                if best_score > 0.7 and best_candidate is not None:
                    conf = self.scorer.calculate_confidence('semantic', best_score, context)
                    return TranslationResult(word, self.dictionary[best_candidate], conf, 'semantic', time.time() - start_time)
            except Exception as e:
                logger.warning(f"Semantic translation failed for '{word}': {e}", exc_info=True)
            # 5. Zero-shot translation as last resort
            try:
                inputs = self.zero_shot_tokenizer(processed, return_tensors="pt", padding=True)
                src_lang_map = {
                    'id': 'id_ID',
                    'dyk': 'id_ID',  # Map Dayak to Indonesian
                    'ms': 'id_ID'
                }
                self.zero_shot_tokenizer.src_lang = src_lang_map.get(self.source_lang, "id_ID")
                target_lang_map = {
                    'id': 'id_ID',
                    'dyk': 'id_ID',
                    'ms': 'id_ID'
                }
                forced_lang = target_lang_map.get(self.target_lang, self.target_lang)
                generated = self.zero_shot_model.generate(
                    **inputs,
                    forced_bos_token_id=self.zero_shot_tokenizer.lang_code_to_id[forced_lang],
                    max_length=128,
                    num_beams=4,
                    length_penalty=0.6,
                    max_time=10.0
                )
                zs_trans = self.zero_shot_tokenizer.batch_decode(generated, skip_special_tokens=True)[0]
                conf = self.scorer.calculate_confidence('zero_shot', 0.5, context)
                return TranslationResult(word, zs_trans, conf, 'zero_shot', time.time() - start_time)
            except Exception as e:
                logger.warning(f"Zero-shot translation failed for '{word}': {e}", exc_info=True)
            # 6. Fallback to original word with very low confidence.
            return TranslationResult(word, word, 0.1, 'fallback', time.time() - start_time)
        finally:
            if torch is not None and torch.cuda.is_available():
                torch.cuda.empty_cache()

    def apply_fst_rules(self, word: str) -> Tuple[Optional[str], float]:
        """
        Apply regex-based FST rules with extended patterns.
        """
        rules = [
            (r'([a-zA-Z]+)s$', r'\1'),
            (r'([a-zA-Z]+)an$', r'\1'),
            (r'([a-zA-Z]+)nya$', r'\1'),
            (r'^(un|in)([a-zA-Z]+)$', r'\2')
        ]
        for pattern, repl in rules:
            candidate = re.sub(pattern, repl, word)
            if candidate in self.dictionary:
                return candidate, 0.85
        return None, 0.0

    def apply_rbmt_rules(self, word: str) -> Tuple[Optional[str], float]:
        """
        Apply Rule-Based Machine Translation (RBMT) rules.
        This includes handling common idiomatic expressions and morphological patterns.
        """
        rbmt_rules = [
            (r'\bselamat pagi\b', 'good morning', 0.9),
            (r'\bapa kabar\b', 'how are you', 0.9)
        ]
        for pattern, replacement, conf in rbmt_rules:
            if re.search(pattern, word):
                return replacement, conf
        if word.endswith("ing"):
            candidate = word[:-3]
            if candidate in self.dictionary:
                return self.dictionary[candidate], 0.8
        return None, 0.0

    def translate_sentence(self, sentence: str) -> str:
        """
        Translate a sentence while preserving word order.
        """
        tokens = sentence.split()
        results = [None] * len(tokens)
        with ThreadPoolExecutor(max_workers=self.resource_limits.thread_count) as executor:
            futures = {executor.submit(self.translate_word, token): i for i, token in enumerate(tokens)}
            try:
                for future in as_completed(futures, timeout=60):
                    idx = futures[future]
                    try:
                        res = future.result(timeout=60)
                        results[idx] = res.translated
                    except Exception as e:
                        logger.error("Token translation error.", exc_info=True)
                        results[idx] = tokens[idx]
            except Exception as e:
                logger.error("Sentence translation timeout or error.", exc_info=True)
        results = [token if token is None else token for token in results]
        return " ".join(results)

    def translate_text(self, text: str, source_lang: str, target_lang: str) -> str:
        logger.info(f"Translating text from {source_lang} to {target_lang}.")
        # Normalize the entire text before splitting into sentences.
        normalized_text = self.text_processor.normalize_text(text)
        sentences = self.text_processor.split_sentences(normalized_text)
        translated_sentences: List[str] = []
        for sentence in sentences:
            word_translations = [self.translate_word(token) for token in sentence.split()]
            avg_conf = np.mean([res.confidence for res in word_translations]) if word_translations else 0.0
            if avg_conf < 0.4:
                try:
                    inputs = self.zero_shot_tokenizer(sentence, return_tensors="pt", padding=True)
                    target_lang_map = {
                        'id': 'id_ID',
                        'dyk': 'id_ID',
                        'ms': 'id_ID'
                    }
                    forced_lang = target_lang_map.get(target_lang, target_lang)
                    self.zero_shot_tokenizer.src_lang = source_lang
                    generated = self.zero_shot_model.generate(
                        **inputs,
                        forced_bos_token_id=self.zero_shot_tokenizer.lang_code_to_id[forced_lang],
                        max_length=128,
                        num_beams=4,
                        length_penalty=0.6
                    )
                    sentence_trans = self.zero_shot_tokenizer.batch_decode(generated, skip_special_tokens=True)[0]
                    translated_sentences.append(sentence_trans)
                    continue
                except Exception as e:
                    logger.warning("Sentence-level zero-shot translation failed.", exc_info=True)
            translated_sentences.append(" ".join([res.translated for res in word_translations]))
        overall_translation = " ".join(translated_sentences)
        gc.collect()
        return overall_translation

    def create_report(self, filename: str) -> None:
        doc = Document()
        doc.add_heading("ILTE-ZS Translation Report", 0)
        total_conf = []
        total_words = 0
        for original, translated, conf, orig_count, _ in getattr(self, 'translations', []):
            doc.add_paragraph(f"Original: {original}")
            doc.add_paragraph(f"Translated: {translated}")
            doc.add_paragraph(f"Confidence: {conf:.1%}")
            doc.add_paragraph("")
            total_conf.append(conf)
            total_words += orig_count
        avg_conf = np.mean(total_conf) if total_conf else 0.0
        footer_text = f"Total Words Processed: {total_words} | Average Confidence: {avg_conf:.1%}"
        section = doc.sections[0]
        if section.footer.paragraphs:
            section.footer.paragraphs[0].text = footer_text
        else:
            section.footer.add_paragraph(footer_text)
        doc.add_page_break()
        doc.save(filename)
        logger.info(f"Report saved to {filename}")

    def cleanup(self) -> None:
        """Cleanup resources: shutdown executors, save cache, and release GPU memory."""
        self.parallel_processor.shutdown()
        self.resource_manager.cache_manager.save()
        self.resource_manager.clean_cache()
        self.resource_manager.release_gpu_memory()
        gc.collect()

class ILTEZSEngineRunner:
    """
    Main interface and execution workflow for the ILTE-ZS engine.
    Supports interactive mode and file input.
    """
    def __init__(self, engine: EnhancedILTEZSTranslationEngine):
        self.engine = engine
        signal.signal(signal.SIGINT, self._handle_interrupt)
        signal.signal(signal.SIGTERM, self._handle_interrupt)

    def _handle_interrupt(self, signum, frame):
        """Handle interrupt signals gracefully."""
        print("\nReceived interrupt signal. Cleaning up...")
        self.engine.shutdown_event.set()
        self.engine.cleanup()
        print("Cleanup complete. Exiting...")
        sys.exit(0)

    def process_text(self, text: str) -> List[str]:
        sentences = self.engine.text_processor.split_sentences(text)
        results: List[str] = []
        with ThreadPoolExecutor(max_workers=self.engine.resource_limits.thread_count) as executor:
            futures = {executor.submit(self.engine.translate_sentence, s): s for s in sentences}
            for future in as_completed(futures, timeout=120):
                try:
                    results.append(future.result(timeout=60))
                except Exception as e:
                    logger.error("Sentence translation error.", exc_info=True)
        return results

    def process_file(self, file_path: str) -> List[str]:
        try:
            safe_path = Path(file_path).resolve()
            if not safe_path.exists():
                logger.error("File not found.")
                return []
            encoding = detect_encoding(str(safe_path))
            if safe_path.suffix.lower() == ".docx":
                doc = Document(str(safe_path))
                text = "\n".join([para.text for para in doc.paragraphs])
            else:
                with safe_path.open('r', encoding=encoding, errors='replace') as file:
                    text = file.read()
            return self.process_text(text)
        except Exception as e:
            logger.error("File processing error.", exc_info=True)
            return []

    def generate_report(self, results: List[str], filename: str = "translation_report.docx") -> None:
        doc = Document()
        doc.add_heading("ILTE-ZS Translation Report", 0)
        for sentence in results:
            doc.add_paragraph(sentence)
        doc.save(filename)
        logger.info(f"Report saved as {filename}")

    def interactive_mode(self) -> None:
        print("\n--- Interactive Translation Mode ---")
        while not self.engine.shutdown_event.is_set():
            try:
                text = input("\nEnter text to translate (or type 'exit' to quit): ").strip()
                if text.lower() == "exit":
                    break
                print("Translating...")
                results = self.process_text(text)
                print("\n--- Translation Results ---")
                for res in results:
                    print(f"→ {res}")
            except KeyboardInterrupt:
                print("\nTranslation interrupted by user.")
                break
            except Exception as e:
                print(f"\nError occurred: {str(e)}")
                logger.error("Interactive mode error", exc_info=True)
                break

    def file_mode(self) -> None:
        file_path = input("\nEnter file path: ").strip()
        if not os.path.exists(file_path):
            print("File not found.")
            return
        print("\nDetecting file encoding...")
        encoding = detect_encoding(file_path)
        print(f"Detected encoding: {encoding}")
        print("Processing file...")
        results = self.process_file(file_path)
        print("\n--- Preview (first 5 sentences) ---")
        for res in results[:5]:
            print(f"→ {res}")
        self.generate_report(results)

    def main(self) -> None:
        print("\n=== Welcome to ILTE-ZS Translation Engine ===")
        while True:
            mode = input("\nChoose mode: [1] Direct Input  [2] File Input  [3] Exit: ").strip()
            if mode == "1":
                self.interactive_mode()
            elif mode == "2":
                self.file_mode()
            elif mode == "3":
                print("Exiting translation engine...")
                break
            else:
                print("Invalid choice. Try again.")

def detect_encoding(file_path: str, read_bytes: int = 100000) -> str:
    """
    Detect file encoding by reading up to `read_bytes` from the file.
    """
    try:
        with open(file_path, "rb") as f:
            raw_data = f.read(read_bytes)
        detected = chardet.detect(raw_data)
        encoding = detected.get("encoding", "utf-8")
        logger.info(f"Detected encoding: {encoding}")
        return encoding
    except Exception as e:
        logger.error("Encoding detection error.", exc_info=True)
        return "utf-8"

if __name__ == "__main__":
    try:
        dictionary_path = "dictionary_alt.json"
        engine = EnhancedILTEZSTranslationEngine(dictionary_path)
        runner = ILTEZSEngineRunner(engine)
        runner.main()
    except Exception as main_exc:
        logger.critical("Fatal error in ILTE-ZS engine execution.", exc_info=True)
    finally:
        engine.cleanup()
        logger.info("ILTE-ZS engine cleanup complete.")
        sys.exit(0)