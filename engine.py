import json
import numpy as np
from docx import Document
from nltk.stem import PorterStemmer, SnowballStemmer
from sentence_transformers import SentenceTransformer
from transformers import pipeline
from Levenshtein import distance as lev_distance

class IndigenousTranslator:
    def __init__(self, json_path, doc_path):
        self.dictionary = self.load_json(json_path)
        self.examples = self.load_doc_examples(doc_path)
        self.embedder = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
        self.example_embeddings = self.precompute_embeddings()
        self.en_stemmer = PorterStemmer()
        self.id_stemmer = SnowballStemmer("indonesian")
        self.confidence = 1.0
        self.match_rate = 0.0

    def load_json(self, path):
        with open(path, 'r') as f:
            return {k.lower(): v for k, v in json.load(f).items()}

    def load_doc_examples(self, path):
        doc = Document(path)
        return [tuple(map(str.strip, para.text.split('||'))) 
                for para in doc.paragraphs if '||' in para.text]

    def precompute_embeddings(self):
        return self.embedder.encode([ex[0] for ex in self.examples])

    def preprocess_text(self, text, source_lang):
        text = text.lower()
        if source_lang == 'en':
            return ' '.join([self.en_stemmer.stem(word) for word in text.split()])
        elif source_lang == 'id':
            return ' '.join([self.id_stemmer.stem(word) for word in text.split()])
        return text

    def find_closest_word(self, word, source_lang):
        candidates = [k for k in self.dictionary.keys() 
                     if source_lang in self.dictionary[k]]
        if not candidates:
            return word, 0.0
        distances = [(c, lev_distance(word, c)) for c in candidates]
        closest = min(distances, key=lambda x: x[1])
        max_len = max(len(word), len(closest[0]))
        return closest[0], 1 - (closest[1]/max_len)

    def translate_word(self, word, source_lang):
        original = word.lower()
        stemmed = self.preprocess_text(original, source_lang)
        
        # Try exact match
        if original in self.dictionary:
            return self.dictionary[original]['indigenous'], 1.0
        
        # Try stemmed match
        if stemmed in self.dictionary:
            self.confidence *= 0.9
            return self.dictionary[stemmed]['indigenous'], 0.9
        
        # Find closest word
        closest, confidence = self.find_closest_word(original, source_lang)
        if confidence > 0.6:
            self.confidence *= confidence
            return self.dictionary[closest]['indigenous'], confidence
        
        return original, 0.0

    def translate_sentence(self, text, source_lang):
        words = text.split()
        total = len(words)
        matched = 0
        translations = []
        example_emb = self.embedder.encode([text])[0]
        example_sim = np.dot(example_emb, self.example_embeddings.T).max()
        
        # Check example-based translation
        if example_sim > 0.8:
            best_match = self.examples[np.argmax(example_sim)]
            self.confidence = example_sim
            self.match_rate = 1.0
            return best_match[1]

        # Word-by-word translation
        for word in words:
            trans, score = self.translate_word(word, source_lang)
            translations.append(trans)
            if score > 0.6:
                matched += 1
                
        self.match_rate = matched / total
        self.confidence = min(self.confidence, 0.9)
        return ' '.join(translations)

    def create_report(self, original, translation, filename):
        doc = Document()
        doc.add_heading('Translation Report', 0)
        
        doc.add_paragraph('Original Text:').bold = True
        doc.add_paragraph(original)
        
        doc.add_paragraph('Translation:').bold = True
        doc.add_paragraph(translation)
        
        doc.add_paragraph(f'Match Rate: {self.match_rate:.1%}').bold = True
        doc.add_paragraph(f'Confidence: {self.confidence:.1%}').bold = True
        
        doc.save(filename)

# Usage Example
if __name__ == "__main__":
    translator = IndigenousTranslator('dictionary.json', 'examples.docx')
    
    input_text = input("Enter the text to translate: \n")
    source_lang = 'id'  # or 'en'
    
    translation = translator.translate_sentence(input_text, source_lang)
    translator.create_report(input_text, translation, 'translation_report.docx')
    
    print(f"Translation: {translation}")
    print(f"Confidence: {translator.confidence:.1%}")
    print(f"Match Rate: {translator.match_rate:.1%}")