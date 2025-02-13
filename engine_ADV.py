import json
import gc
import numpy as np
import chardet
import re
from Levenshtein import distance as lev_distance
from nltk.stem import SnowballStemmer
from Sastrawi.Stemmer.StemmerFactory import StemmerFactory
from transformers import (
    pipeline,
    AutoTokenizer,
    AutoModelForMaskedLM
)
from sentence_transformers import SentenceTransformer
import torch
from torch.utils.data import DataLoader, Dataset
from docx import Document
import os
from collections import defaultdict
from tqdm import tqdm
from typing import List, Tuple, Dict

def sanitize_text(text: str) -> str:
    """
    Remove control characters and other XML-incompatible characters.
    """
    # Remove characters in ranges: U+0000-U+0008, U+000B-U+000C, U+000E-U+001F
    return re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F]', '', text)

class TextBatchDataset(Dataset):
    def __init__(self, texts: List[str], window_size: int = 3):
        self.texts = texts
        self.window_size = window_size
        self.windows = self._create_windows()
        
    def _create_windows(self):
        windows = []
        for text in self.texts:
            words = text.split()
            for i in range(len(words)):
                window = words[max(0, i-self.window_size):min(len(words), i+self.window_size+1)]
                windows.append((" ".join(window), words[i]))
        return windows
    
    def __len__(self):
        return len(self.windows)
        
    def __getitem__(self, idx):
        return self.windows[idx]

class OptimizedIndigenousTranslator:
    def __init__(self, json_path, batch_size=32):
        # Load dictionary and initialize stemmers
        self.dictionary = self.load_json(json_path)
        self.batch_size = batch_size
        factory = StemmerFactory()
        self.id_stemmer = factory.create_stemmer()
        self.en_stemmer = SnowballStemmer("english")
        
        # Optimization config for translation pipelines
        self.optimization_config_translation = {
            "torch_dtype": torch.float16,
            "load_in_8bit": True,
            "max_memory": {"cuda:0": "3.5GB"}
        }
        
        # General optimization config for other models (excluding zero-shot classifier)
        self.optimization_config = {
            "torch_dtype": torch.float16,
            "device_map": "auto",
            "load_in_8bit": True,
            "max_memory": {"cuda:0": "3.5GB"}
        }
        
        # Separate optimization config for zero-shot classifier (without device_map)
        self.optimization_config_zero_shot = {
            "torch_dtype": torch.float16,
            "load_in_8bit": True,
            "max_memory": {"cuda:0": "3.5GB"}
        }
        
        # Initialize translation pipelines
        print("Loading translation models...")
        self.en_id_translator = pipeline(
            "translation_en_to_id",
            model="Helsinki-NLP/opus-mt-en-id",
            tokenizer="Helsinki-NLP/opus-mt-en-id",
            **self.optimization_config_translation
        )
        self.id_en_translator = pipeline(
            "translation_id_to_en",
            model="Helsinki-NLP/opus-mt-id-en",
            tokenizer="Helsinki-NLP/opus-mt-id-en",
            **self.optimization_config_translation
        )
        
        # Load Indonesian BERT without extra device_map (use as is)
        print("Loading BERT model...")
        self.tokenizer = AutoTokenizer.from_pretrained("cahya/bert-base-indonesian-1.5G")
        self.context_model = AutoModelForMaskedLM.from_pretrained(
            "cahya/bert-base-indonesian-1.5G",
            torch_dtype=torch.float16
        )
        if torch.cuda.is_available():
            self.context_model.to("cuda")
        
        # Initialize SentenceTransformer
        print("Loading SentenceTransformer...")
        self.sentence_model = SentenceTransformer(
            'sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2',
            device='cuda' if torch.cuda.is_available() else 'cpu'
        )
        
        # Precompute dictionary embeddings
        print("Computing dictionary embeddings...")
        self.dict_words = list(self.dictionary.keys())
        if self.dict_words:
            self.dict_embeddings = self.sentence_model.encode(
                self.dict_words,
                batch_size=self.batch_size,
                convert_to_numpy=True,
                show_progress_bar=True
            )
        else:
            self.dict_embeddings = np.array([])
        
        # Initialize zero-shot classifier
        print("Loading zero-shot classifier...")
        self.zero_shot = pipeline(
            "zero-shot-classification",
            model="typeform/distilbert-base-uncased-mnli",
            tokenizer="typeform/distilbert-base-uncased-mnli",
            **self.optimization_config_zero_shot
        )
        
        # Translation tracking
        self.translations = []
        self.pattern_memory = defaultdict(dict)
        self.context_cache = {}
        print("Initialization complete!")
        
    @staticmethod
    def detect_encoding(file_path):
        """Detect file encoding using chardet."""
        with open(file_path, "rb") as f:
            raw_data = f.read(1024)
        result = chardet.detect(raw_data)
        return result.get("encoding")
        
    def load_json(self, path):
        with open(path, "r") as f:
            return json.load(f)
            
    @torch.no_grad()
    def batch_encode_contexts(self, contexts: List[str]) -> torch.Tensor:
        """Encode multiple contexts in a single batch."""
        inputs = self.tokenizer(contexts, return_tensors="pt", padding=True, truncation=True)
        inputs = {k: v.to(self.context_model.device) for k, v in inputs.items()}
        outputs = self.context_model(**inputs, output_hidden_states=True)
        return outputs.hidden_states[-1].mean(dim=1)
    
    @torch.no_grad()
    def batch_semantic_similarity(self, words: List[str], contexts: List[str]) -> List[Tuple[str, float]]:
        """Compute semantic similarity for multiple words and contexts in parallel."""
        # Encode words and contexts
        word_embeddings = self.sentence_model.encode(words, batch_size=self.batch_size, convert_to_numpy=True)
        context_embeddings = self.sentence_model.encode(contexts, batch_size=self.batch_size, convert_to_numpy=True)
        # Normalize embeddings
        word_embeddings = word_embeddings / np.linalg.norm(word_embeddings, axis=1, keepdims=True)
        context_embeddings = context_embeddings / np.linalg.norm(context_embeddings, axis=1, keepdims=True)
        # Compute cosine similarities
        word_similarities = np.dot(word_embeddings, self.dict_embeddings.T)
        context_similarities = np.dot(context_embeddings, self.dict_embeddings.T)
        # Combine scores
        combined_scores = 0.7 * word_similarities + 0.3 * context_similarities
        best_indices = np.argmax(combined_scores, axis=1)
        best_scores = combined_scores[np.arange(len(best_indices)), best_indices]
        return [(self.dict_words[idx], score) for idx, score in zip(best_indices, best_scores)]
    
    def analyze_patterns_batch(self, texts: List[str], window_size=3) -> Dict[str, Dict]:
        """Analyze text patterns in batch."""
        patterns = defaultdict(dict)
        dataset = TextBatchDataset(texts, window_size)
        dataloader = DataLoader(dataset, batch_size=self.batch_size, shuffle=False)
        for batch_contexts, batch_words in dataloader:
            context_embeddings = self.batch_encode_contexts(batch_contexts)
            for context, word, embedding in zip(batch_contexts, batch_words, context_embeddings):
                pattern_key = tuple(context.split())
                if pattern_key in self.pattern_memory:
                    patterns[word].update(self.pattern_memory[pattern_key])
        return patterns
    
    def preprocess_text(self, text: str, lang: str) -> Tuple[str, Dict[str, str]]:
        """Enhanced preprocessing with context preservation."""
        text = text.lower()
        words = text.split()
        if lang == "id":
            processed = [self.id_stemmer.stem(word) for word in words]
        elif lang == "en":
            processed = [self.en_stemmer.stem(word) for word in words]
        else:
            processed = words
        mapping = dict(zip(words, processed))
        return " ".join(processed), mapping
    
    def translate_text(self, text: str, source_lang: str, target_lang: str) -> str:
        """Optimized text translation using batched processing."""
        print(f"Translating from {source_lang} to {target_lang}...")
        
        if source_lang == "en" and target_lang == "dyk":
            text = self.en_id_translator(text, max_length=512)[0]["translation_text"]
        elif source_lang == "dyk" and target_lang == "en":
            text = self.translate_text(text, "dyk", "id")
            text = self.id_en_translator(text, max_length=512)[0]["translation_text"]
            return text
            
        processed_text, word_mapping = self.preprocess_text(text, source_lang)
        words = processed_text.split()
        
        # Create dataset and dataloader for batch processing
        dataset = TextBatchDataset(words)
        dataloader = DataLoader(dataset, batch_size=self.batch_size, shuffle=False)
        
        translated_words = []
        confidence_list = []
        
        print("Processing translation in batches...")
        for batch_contexts, batch_words in tqdm(dataloader, desc="Translating"):
            semantic_matches = self.batch_semantic_similarity(batch_words, batch_contexts)
            for word, (match, confidence) in zip(batch_words, semantic_matches):
                if confidence > 0.8:
                    translation = self.dictionary.get(match, word)
                    translated_words.append(translation)
                    confidence_list.append(confidence)
                else:
                    translation = self.dictionary.get(word, word)
                    translated_words.append(translation)
                    confidence_list.append(1.0 if word in self.dictionary else 0.0)
        
        self.confidence = np.mean(confidence_list)
        self.match_rate = len([c for c in confidence_list if c > 0]) / len(confidence_list)
        
        translation = " ".join(translated_words)
        self.translations.append((
            text,
            translation,
            self.confidence,
            len(words),
            len(translated_words)
        ))
        
        print(f"Translation complete! Confidence: {self.confidence:.2%}")
        return translation
    
    def create_report(self, filename: str):
        """Enhanced report creation with additional metrics."""
        doc = Document()
        doc.add_heading("Translation Report", 0)
        
        avg_confidence = sum(conf for _, _, conf, _, _ in self.translations) / len(self.translations)
        total_words = sum(orig_count for _, _, _, orig_count, _ in self.translations)
        total_translations = sum(trans_count for _, _, _, _, trans_count in self.translations)
        
        header = doc.sections[0].header
        header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_paragraph.text = (
            f"Performance Score: {avg_confidence:.1%} | "
            f"Translation Rate: {self.match_rate:.1%} | "
            f"Pattern Matches: {len(self.pattern_memory)}"
        )
        
        for original, translated, confidence, original_word_count, target_word_count in self.translations:
            doc.add_paragraph(f"Original: {sanitize_text(original)}")
            doc.add_paragraph(f"Translated: {sanitize_text(translated)}")
            doc.add_paragraph(f"Confidence: {confidence:.1%}")
            doc.add_paragraph("\n")
        
        footer = doc.sections[0].footer
        footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_paragraph.text = (
            f"Total Words Processed: {total_words} | "
            f"Total Translations: {total_translations} | "
            f"Average Confidence: {avg_confidence:.1%}"
        )
        
        doc.add_page_break()
        doc.save(filename)
        print(f"Report saved to {filename}")

def main():
    try:
        translator = OptimizedIndigenousTranslator("dictionary_alt.json")
        
        while True:
            choice = input("\n\nEnter 'file' to load from file or 'text' to input manually: ")
            
            if choice.lower() == 'file':
                file_path = input("Enter file path: ")
                if os.path.exists(file_path):
                    encoding = OptimizedIndigenousTranslator.detect_encoding(file_path)
                    if not encoding:
                        print("Encoding could not be detected. Trying fallback encoding 'latin-1'...")
                        encoding = "latin-1"
                    try:
                        with open(file_path, "r", encoding=encoding) as file:
                            text = file.read()
                    except UnicodeDecodeError:
                        print(f"Error decoding file with {encoding}. Trying fallback encoding 'latin-1'...")
                        with open(file_path, "r", encoding="latin-1") as file:
                            text = file.read()
                else:
                    print("File not found.")
                    continue
            else:
                text = input("Enter text to translate: ")
            
            source_lang = input("Enter source language (id/en/dyk): ").lower()
            target_lang = input("Enter target language (id/en/dyk): ").lower()
            print("-" * 50)
            print("\n")
            
            if source_lang not in ['id', 'en', 'dyk'] or target_lang not in ['id', 'en', 'dyk']:
                print("Invalid language selection. Please try again.")
                continue
            
            translation = translator.translate_text(text, source_lang, target_lang)
            translator.create_report("translation_report.docx")
            
            print("\nTranslation Results:")
            print("-" * 50)
            print(f"Translation: {translation}")
            print(f"Confidence: {translator.confidence:.1%}")
            print(f"Match Rate: {translator.match_rate:.1%}")
            print("-" * 50)
            
            continue_translation = input("\nWould you like to translate another text? (y/n): ")
            if continue_translation.lower() != 'y':
                break
                
    except KeyboardInterrupt:
        print("\nTranslation process interrupted.")
    finally:
        print("\nCleaning up resources...")
        torch.cuda.empty_cache()
        gc.collect()
        print("Goodbye!")

if __name__ == "__main__":
    main()